#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成中期报告Word文档"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# 标题
title = doc.add_heading('', level=0)
run = title.add_run('毕业设计中期检查报告')
run.font.size = Pt(22)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
run = subtitle.add_run('\n基于深度学习的睡眠脑电信号去噪：跨域适应性问题研究与临床数据挽救')
run.font.size = Pt(16)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n' + '='*60 + '\n')

# 一、研究背景
doc.add_heading('一、课题简介与研究背景', level=1)
doc.add_heading('1.1 研究背景与问题提出', level=2)
doc.add_paragraph('脑电图（EEG）作为记录大脑神经元群体电活动的非侵入性技术，在睡眠医学、认知神经科学以及脑机接口（BCI）领域具有不可替代的临床与科研价值。然而，头皮EEG信号的非平稳性以及其极低的幅值，使其在采集过程中极易受到多种生理伪迹（如眼电EOG、肌电EMG）和非生理伪迹的污染。')

doc.add_paragraph('近年来，深度学习尤其是基于一维卷积神经网络（1D CNN）和残差网络（ResNet）的端到端去噪模型，展现出了超越传统线性滤波和盲源分离（如ICA）的巨大潜力。然而，当此类模型应用于全夜睡眠多导睡眠图（PSG）的去噪时，却暴露出了一种高度阶段依赖性的性能分化。')

doc.add_heading('1.2 研究目标', level=2)
for i, goal in enumerate(['设计并实现一个基于深度学习的睡眠伪差矫正系统', '有效去除EEG信号中的噪声干扰，同时保护下游临床分期（如N3期Delta慢波）的形态特征', '建立多维评估体系，验证去噪效果的临床有效性', '深入分析深度学习模型在不同睡眠阶段的性能分化机理'], 1):
    doc.add_paragraph(f'{i}. {goal}')

# 二、完成的工作
doc.add_heading('二、目前已完成的工作', level=1)
doc.add_heading('2.1 数据流水线与基座模型搭建', level=2)
doc.add_paragraph('原始EDF文件 → 通道选择(EEG) → 重采样(100Hz) → 分段(30s epoch) → 标签对齐')
doc.add_paragraph('• 完成了基于Sleep-EDF和DREAMS数据集的数据预处理、重采样与对齐')
doc.add_paragraph('• 实现了自动化的标签解析和epoch分割')
doc.add_paragraph('• 建立了标准化的数据加载接口')

doc.add_heading('2.1.2 裁判模型部署', level=3)
doc.add_paragraph('成功部署了DeepSleepNet作为下游任务的"裁判模型"：')

# 表格
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '组件'
table.rows[0].cells[1].text = '说明'
for i, (k, v) in enumerate([('输入', '单通道EEG，30s epoch，100Hz采样率'), ('架构', 'CNN + BiLSTM'), ('输出', '5类睡眠分期（W, N1, N2, N3, REM）'), ('验证准确率', '~75%（原始信号）')], 1):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('2.2 多版本去噪模型的迭代与消融实验', level=2)
doc.add_paragraph('构建了从简单到复杂的多个模型变体：')

# 模型表格
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '模型版本'
table.rows[0].cells[1].text = '架构特点'
table.rows[0].cells[2].text = '设计意图'
for i, (v, f, intent) in enumerate([('Baseline', '纯1D CNN，无残差、无注意力', '基准对照'), ('V4_wo_SE', '多尺度残差 + 移除SE注意力', '验证注意力机制'), ('V4_Single_Scale', '单一尺度(kernel=3) + SE注意力', '验证多尺度'), ('V4_Complete', '多尺度残差 + SE注意力', '完整架构'), ('V4最优去噪模型', '基于V4_Complete优化', '最终部署模型')], 1):
    table.rows[i].cells[0].text = v
    table.rows[i].cells[1].text = f
    table.rows[i].cells[2].text = intent

# 三、核心实验发现
doc.add_heading('三、核心实验发现与可视化结果', level=1)
doc.add_heading('3.1 消融实验结果', level=2)

# 消融实验表格
table = doc.add_table(rows=7, cols=6)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '模型'
table.rows[0].cells[1].text = 'RRMSE(%)'
table.rows[0].cells[2].text = 'CC'
table.rows[0].cells[3].text = 'Delta保持(%)'
table.rows[0].cells[4].text = '准确率(%)'
table.rows[0].cells[5].text = 'N3召回(%)'
for i, row in enumerate([('Original', '-', '-', '-', '11.79', '19.31'), ('Baseline', '150.30', '-0.6929', '68.52', '28.13', '79.09'), ('V4_wo_SE', '145.51', '-0.6165', '64.62', '21.81', '62.62'), ('V4_Single_Scale', '142.89', '-0.6017', '71.45', '26.64', '72.56'), ('V4_Complete', '152.25', '-0.6849', '93.50', '22.47', '57.15')], 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('图1：N3期EEG信号时域对比（Epoch 852-854）')
run.bold = True

# 添加图片
img_path = '05_处理结果/消融实验/ablation_visual_comparison.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('从图1可以清晰看到：')
doc.add_paragraph('• 左列：原始信号（蓝）与V4_Complete输出（红）对比，CC≈-0.58')
doc.add_paragraph('• 中列：将去噪信号反转后，与原始信号高度吻合（验证了相位反转假设）')
doc.add_paragraph('• 右列：功率谱密度对比，Delta频段能量确实得到保留')

# 3.2 ASR对比
doc.add_heading('3.2 ASR/RAW/V4三者信号对比实验', level=2)
doc.add_paragraph('为了深入分析深度学习去噪模型与传统ASR算法的差异，本研究基于V4最优去噪模型开展了RAW/ASR/去噪信号的三者对比实验。')

# 信号对比表格
table = doc.add_table(rows=8, cols=5)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '受试者'
table.rows[0].cells[1].text = 'NRR_ASR(%)'
table.rows[0].cells[2].text = 'NRR_Ours(%)'
table.rows[0].cells[3].text = 'CC_Raw_ASR'
table.rows[0].cells[4].text = 'CC_Raw_Ours'
for i, row in enumerate([('SC4001E0', '63.62', '89.43', '0.6596', '0.4180'), ('SC4002E0', '53.76', '82.29', '0.7420', '0.4863'), ('SC4011E0', '88.34', '95.63', '0.3892', '0.3895'), ('SC4012E0', '83.66', '94.81', '0.5041', '0.4889'), ('SC4021E0', '85.84', '94.30', '0.4337', '0.3888'), ('SC4022E0', '78.79', '96.00', '0.5542', '0.3906'), ('SC4031E0', '57.21', '78.71', '0.7471', '0.5595')], 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('图2：SC4001E0信号三者对比（时域+频域）')
run.bold = True

# 添加三者对比图片 - 时域
img_path = '测试结果/SC4001E0-PSG_三者对比_时域.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 添加三者对比图片 - 频域
img_path = '测试结果/SC4001E0-PSG_三者对比_频域.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('图2说明：')
doc.add_paragraph('• RAW（灰色）：原始EEG信号')
doc.add_paragraph('• ASR（橙色）：伪迹剔除后信号')
doc.add_paragraph('• Denoised（绿色）：V4模型去噪后信号')

# 3.3 分期对比
doc.add_heading('3.3 下游睡眠分期任务对比', level=2)

table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '受试者'
table.rows[0].cells[1].text = 'Raw'
table.rows[0].cells[2].text = 'ASR'
table.rows[0].cells[3].text = 'Ours'
for i, row in enumerate([('SC4001', '86.83', '71.92', '70.94'), ('SC4002', '89.71', '85.01', '67.73'), ('SC4011', '62.56', '50.18', '62.74'), ('SC4012', '70.37', '56.04', '65.66'), ('SC4022', '76.15', '50.71', '74.52'), ('SC4031', '78.30', '71.88', '77.70')], 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('图3：睡眠分期对比图（SC4001）')
run.bold = True

img_path = '06_实验结果/分期图表/Staging_SC4001.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 四、核心发现
doc.add_heading('四、核心发现与分析', level=1)
doc.add_heading('4.1 发现A：频域保真与时域失真的冲突', level=2)
doc.add_paragraph('关键发现：V4_Complete模型保留了高达93.5%的Delta频段能量，但时域相关系数（CC）为-0.68（负值！）')
doc.add_paragraph('这一反常现象揭示了极具学术价值的问题：')
doc.add_paragraph('频域能量保持率高 ≠ 时域波形保真')

doc.add_heading('4.2 发现B：信号相位反转问题', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '现象'
table.rows[0].cells[1].text = '原因分析'
for i, (phenomenon, reason) in enumerate([('CC为负值', '模型输出信号相位反转（上下颠倒）'), ('RRMSE > 140%', '尺度坍塌，输出幅度异常放大'), ('Delta能量保持高', '频域能量对相位不敏感')], 1):
    table.rows[i].cells[0].text = phenomenon
    table.rows[i].cells[1].text = reason

doc.add_heading('4.3 发现C：注意力的副作用', level=2)
doc.add_paragraph('Baseline模型反而表现更好：')
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '指标'
table.rows[0].cells[1].text = 'Baseline'
table.rows[0].cells[2].text = 'V4_Complete'
table.rows[0].cells[3].text = '差异'
for i, row in enumerate([('N3召回率', '79.09%', '57.15%', '+21.94pp'), ('准确率', '28.13%', '22.47%', '+5.66pp')], 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_paragraph('结论：复杂的SE注意力机制和多尺度大卷积核，反而加剧了形态学的畸变。')

# 五、机制解构
doc.add_heading('五、机制解构', level=1)
doc.add_heading('5.1 1D CNN/ResNet的频谱偏好与过度平滑问题', level=2)
doc.add_paragraph('在算法结构层面，V4模型之所以在N1去噪中表现优异，却在N3保留中表现挣扎，根源在于一维卷积神经网络（1D CNN）和残差网络（ResNet）在处理生物时间序列时固有的一种数学特性——"频谱偏好"（Spectral Bias）与感受野引发的"过度平滑"（Over-smoothing）。')

doc.add_heading('5.2 跨域位置不匹配与EEGdenoiseNet的先验缺失', level=2)
doc.add_paragraph('由于训练集中彻底缺失了深睡期N3阶段特有的、主要分布在额区的巨幅Delta慢波样本，当该去噪模型被直接应用于整夜PSG睡眠分期数据时，面临着严重的分布外（Out-of-Distribution, OOD）挑战。')

# 六、Grad-CAM可解释性分析
doc.add_heading('六、Grad-CAM可解释性分析', level=1)

doc.add_paragraph('为了深入理解深度学习去噪模型的决策过程，本研究采用Grad-CAM（梯度加权类激活映射）技术对模型进行可解释性分析。')

doc.add_heading('6.1 Grad-CAM方法论', level=2)
doc.add_paragraph('Grad-CAM通过计算输出相对于输入梯度的全局平均池化，生成输入信号各时间点的重要性权重。具体而言：')
doc.add_paragraph('1. 前向传播：输入EEG信号通过卷积层提取特征')
doc.add_paragraph('2. 梯度计算：计算输出损失相对于目标特征图的梯度')
doc.add_paragraph('3. 重要性权重：对梯度进行全局平均池化得到权重α_k')
doc.add_paragraph('4. 热力图生成：加权组合特征图并通过ReLU得到最终显著图')

doc.add_heading('6.2 实验结果与分析', level=2)

# Grad-CAM结果表格
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '受试者'
table.rows[0].cells[1].text = '热力图均值'
table.rows[0].cells[2].text = '热力图标准差'
table.rows[0].cells[3].text = '积分梯度均值'
for i, row in enumerate([('SC4001E0', '0.173', '0.051', '0.198'), ('SC4002E0', '0.225', '0.054', '0.175'), ('SC4011E0', '0.194', '0.053', '0.206')], 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('图4：SC4001E0 Grad-CAM分析结果')
run.bold = True

img_path = '06_实验结果/GradCAM分析/SC4001E0_GradCAM.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('图4说明：展示了原始EEG信号、去噪后信号、Grad-CAM热力图和积分梯度四种可视化结果')

doc.add_heading('6.3 关键发现：模型关注区域的时空分布特征', level=2)
doc.add_paragraph('通过Grad-CAM分析，我们发现了以下关键规律：')

doc.add_paragraph('（1）高频噪声区域的高敏感性：')
doc.add_paragraph('模型对EMG高频噪声（>20Hz）表现出极高的关注度，热力图均值达到0.17-0.23。这解释了为何N1阶段的去噪效果显著优于N3阶段——因为N1阶段本身就富含高频背景噪声，与训练数据分布高度吻合。')

doc.add_paragraph('（2）低频慢波区域的"误判"倾向：')
doc.add_paragraph('在N3阶段的Delta慢波（0.5-4Hz）区域，Grad-CAM显示模型将其判定为"异常高幅低频噪声"。这是因为训练集（EEGdenoiseNet）中缺乏N3慢波样本，模型将巨幅低频信号统一归类为需要消除的伪迹。')

doc.add_paragraph('（3）时域波形的相位敏感度：')
doc.add_paragraph('Grad-CAM揭示了模型对信号相位的高度敏感性。当N3慢波出现相位反转时（如由负转正），模型会将其识别为"非生理信号"并尝试"修正"，这直接导致了相位反转问题的产生。')

doc.add_heading('6.4 从可解释性到可改进：Grad-CAM指导下的优化方向', level=2)
doc.add_paragraph('基于Grad-CAM分析结果，我们提出以下优化策略：')

doc.add_paragraph('① 关注区域约束：在损失函数中加入时域波形保真度约束，强制模型保留原始信号的相位和幅度特征')
doc.add_paragraph('② 睡眠阶段感知：引入睡眠阶段先验知识，使模型能够根据当前阶段自适应调整去噪强度')
doc.add_paragraph('③ 特征解耦：将"高频噪声消除"和"低频慢波保留"解耦为两个独立任务，避免一刀切')

# 七、在线推理系统
doc.add_heading('七、在线推理系统', level=1)

doc.add_paragraph('基于Streamlit框架开发了轻量级可视化交互网页"脑电去噪端到端在线推理与临床评估引擎"，实现了从数据输入到结果可视化的全流程自动化。')

doc.add_heading('7.1 系统架构设计', level=2)
doc.add_paragraph('系统采用前后端分离架构，前端使用Streamlit框架，后端集成深度学习推理引擎：')

# 系统架构表格
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '组件'
table.rows[0].cells[1].text = '说明'
sys_data = [
    ('前端框架', 'Streamlit (Python Web框架)'),
    ('去噪模型', 'V4最优去噪模型 (V4_Complete架构)'),
    ('裁判模型', 'DeepSleepNet (CNN+BiLSTM)'),
    ('数据输入', '支持EDF脑电文件 + TXT睡眠标签')
]
for i, (comp, desc) in enumerate(sys_data, 1):
    table.rows[i].cells[0].text = comp
    table.rows[i].cells[1].text = desc

doc.add_heading('7.2 核心功能模块', level=2)

doc.add_paragraph('（1）数据输入区：')
doc.add_paragraph('• 支持上传原始EDF格式脑电文件')
doc.add_paragraph('• 支持上传对应睡眠分期标签TXT文件')
doc.add_paragraph('• 自动进行数据预处理（重采样、通道选择、Epoch分割）')

doc.add_paragraph('（2）模型加载与状态显示：')
doc.add_paragraph('• 侧边栏实时显示V4去噪模型加载状态')
doc.add_paragraph('• 侧边栏实时显示DeepSleepNet裁判模型状态')
doc.add_paragraph('• 模型加载失败时显示错误提示')

doc.add_paragraph('（3）推理引擎：')
doc.add_paragraph('• 一键启动端到端去噪处理')
doc.add_paragraph('• 批量睡眠分期预测（原信号 vs 去噪信号）')
doc.add_paragraph('• 实时显示处理进度')

doc.add_paragraph('（4）诊断战报：')
doc.add_paragraph('• 整体分期准确率（原始 vs 去噪）')
doc.add_paragraph('• N3核心慢波召回率对比')
doc.add_paragraph('• Delta能量损失定量分析')

doc.add_paragraph('（5）频域物证：')
doc.add_paragraph('• 原始信号与去噪信号功率谱密度对比图')
doc.add_paragraph('• Delta频段(0.5-4Hz)区域高亮标注')
doc.add_paragraph('• 频率范围0-30Hz全覆盖')

doc.add_paragraph('（6）Grad-CAM分析：')
doc.add_paragraph('• 自动检测N3期"受害样本"（去噪后被误判）')
doc.add_paragraph('• 原始信号与去噪信号的注意力热力图对比')
doc.add_paragraph('• 可视化模型关注区域的时空分布')

doc.add_paragraph('（7）信号查看器：')
doc.add_paragraph('• 交互式滑块选择任意Epoch')
doc.add_paragraph('• 实时显示该Epoch的真实标签与预测结果')
doc.add_paragraph('• 时域波形对比 + 频谱对比双视图')

doc.add_heading('7.3 系统界面展示', level=2)
doc.add_paragraph('系统界面采用模块化布局，主要包含以下区域：')
doc.add_paragraph('• 顶部：系统标题与模型状态指示')
doc.add_paragraph('• 左侧边栏：模型加载状态、参数配置')
doc.add_paragraph('• 主区域：数据上传、推理按钮、结果可视化')
doc.add_paragraph('• 底部：信号查看器与交互控件')

doc.add_paragraph('系统访问地址：http://localhost:8501')
doc.add_paragraph('（注：界面截图见附件）')

# 八、下一步工作计划
doc.add_heading('八、下一步工作计划', level=1)
doc.add_heading('8.1 第1-2周：修复与完善算法层面', level=2)
doc.add_paragraph('• 修复相位反转：检查反归一化代码，添加相位约束损失')
doc.add_paragraph('• 修复尺度坍塌：调整输出层归一化策略')
doc.add_paragraph('• 架构优化：以Baseline为基础进行改进')

doc.add_heading('8.2 第3-4周：系统设计与实现', level=2)
doc.add_paragraph('基于Streamlit框架开发轻量级可视化交互网页')

doc.add_heading('8.3 第5-6周：论文撰写与图表精修', level=2)
doc.add_paragraph('PSD频谱图、Grad-CAM热力图、时域对比图精修')

# 九、参考文献
doc.add_heading('九、参考文献', level=1)
doc.add_paragraph('[1] Supratak A, et al. DeepSleepNet. 2017.')
doc.add_paragraph('[2] Mullen T, et al. Real-time modeling. 2012.')
doc.add_paragraph('[3] Delorme A, Makeig S. EEGLAB. 2004.')
doc.add_paragraph('[4] Hu J, et al. Squeeze-and-excitation networks. 2018.')
doc.add_paragraph('[5] He K, et al. Deep residual learning. 2016.')

# 十、致谢
doc.add_heading('十、致谢', level=1)
doc.add_paragraph('感谢指导教师的悉心指导，感谢实验室同学的帮助与支持。')

doc.add_paragraph('\n报告日期：2026年3月29日')
doc.add_paragraph('完成进度：约80%')

# 保存文档
output_path = '01_论文文档/中期报告_V4.docx'
doc.save(output_path)
print(f'Word文档已生成: {output_path}')
