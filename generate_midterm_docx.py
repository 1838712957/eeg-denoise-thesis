#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
生成中期报告Word文档
包含所有图片和详细解释
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# 设置中文字体
def set_chinese_font(run, font_name='微软雅黑', font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# 创建文档
doc = Document()

# 设置文档默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ==================== 标题 ====================
title = doc.add_heading('', level=0)
run = title.add_run('毕业设计中期检查报告')
set_chinese_font(run, '黑体', 22)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
run = subtitle.add_run('\n基于深度学习的睡眠脑电信号去噪：跨域适应性问题研究与临床数据挽救')
set_chinese_font(run, '黑体', 16)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n' + '='*60 + '\n')

# ==================== 一、课题简介与研究背景 ====================
doc.add_heading('一、课题简介与研究背景', level=1)

doc.add_heading('1.1 研究背景与问题提出', level=2)
p = doc.add_paragraph()
p.add_run('脑电图（EEG）作为记录大脑神经元群体电活动的非侵入性技术，在睡眠医学、认知神经科学以及脑机接口（BCI）领域具有不可替代的临床与科研价值。然而，头皮EEG信号的非平稳性以及其极低的幅值，使其在采集过程中极易受到多种生理伪迹（如眼电EOG、肌电EMG）和非生理伪迹的污染。为了从复杂的背景噪声中提取纯净的神经信号，信号去噪成为EEG数据分析链路中最为关键的预处理步骤。')

p = doc.add_paragraph()
p.add_run('近年来，深度学习尤其是基于一维卷积神经网络（1D CNN）和残差网络（ResNet）的端到端去噪模型，展现出了超越传统线性滤波和盲源分离（如ICA）的巨大潜力。然而，当此类模型应用于全夜睡眠多导睡眠图（PSG）的去噪时，却暴露出了一种高度阶段依赖性的性能分化。')

doc.add_heading('1.2 研究目标', level=2)
goals = [
    '设计并实现一个基于深度学习的睡眠伪差矫正系统',
    '有效去除EEG信号中的噪声干扰，同时保护下游临床分期（如N3期Delta慢波）的形态特征',
    '建立多维评估体系，验证去噪效果的临床有效性',
    '深入分析深度学习模型在不同睡眠阶段的性能分化机理'
]
for i, goal in enumerate(goals, 1):
    p = doc.add_paragraph(f'{i}. {goal}')

# ==================== 二、目前已完成的工作 ====================
doc.add_heading('二、目前已完成的工作', level=1)

doc.add_heading('2.1 数据流水线与基座模型搭建', level=2)

doc.add_heading('2.1.1 数据处理流程', level=3)
p = doc.add_paragraph()
p.add_run('原始EDF文件 → 通道选择(EEG) → 重采样(100Hz/256Hz) → 分段(30s epoch) → 标签对齐')

doc.add_paragraph('• 完成了基于Sleep-EDF和DREAMS数据集的数据预处理、重采样与对齐')
doc.add_paragraph('• 实现了自动化的标签解析和epoch分割')
doc.add_paragraph('• 建立了标准化的数据加载接口')

doc.add_heading('2.1.2 裁判模型部署', level=3)
p = doc.add_paragraph()
p.add_run('成功部署了DeepSleepNet作为下游任务的"裁判模型"：')

# 表格：裁判模型
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
cells = table.rows[0].cells
cells[0].text = '组件'
cells[1].text = '说明'
data = [
    ('输入', '单通道EEG，30s epoch，100Hz采样率'),
    ('架构', 'CNN + BiLSTM'),
    ('输出', '5类睡眠分期（W, N1, N2, N3, REM）'),
    ('验证准确率', '~75%（原始信号）')
]
for i, (key, value) in enumerate(data, 1):
    cells = table.rows[i].cells
    cells[0].text = key
    cells[1].text = value

doc.add_heading('2.2 多版本去噪模型的迭代与消融实验', level=2)

doc.add_heading('2.2.1 模型架构演进', level=3)
p = doc.add_paragraph()
p.add_run('构建了从简单到复杂的多个模型变体：')

# 表格：模型演进
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
cells = table.rows[0].cells
cells[0].text = '模型版本'
cells[1].text = '架构特点'
cells[2].text = '设计意图'
model_data = [
    ('Baseline', '纯1D CNN，无残差、无注意力', '基准对照'),
    ('V4_wo_SE', '多尺度残差 + 移除SE注意力', '验证注意力机制的作用'),
    ('V4_Single_Scale', '单一尺度(kernel=3) + SE注意力', '验证多尺度的必要性'),
    ('V4_Complete', '多尺度残差 + SE注意力', '完整架构'),
    ('V4最优去噪模型', '基于V4_Complete优化', '最终部署模型')
]
for i, (version, feature, intent) in enumerate(model_data, 1):
    cells = table.rows[i].cells
    cells[0].text = version
    cells[1].text = feature
    cells[2].text = intent

# ==================== 三、核心实验发现与可视化结果 ====================
doc.add_heading('三、核心实验发现与可视化结果', level=1)

doc.add_heading('3.1 消融实验结果', level=2)

# 表格：消融实验
table = doc.add_table(rows=7, cols=6)
table.style = 'Table Grid'
cells = table.rows[0].cells
cells[0].text = '模型'
cells[1].text = 'RRMSE(%)'
cells[2].text = 'CC'
cells[3].text = 'Delta保持(%)'
cells[4].text = '准确率(%)'
cells[5].text = 'N3召回(%)'

ablation_data = [
    ('Original (基准)', '-', '-', '-', '11.79', '19.31'),
    ('Baseline', '150.30', '-0.6929', '68.52', '28.13', '79.09'),
    ('V4_wo_SE', '145.51', '-0.6165', '64.62', '21.81', '62.62'),
    ('V4_Single_Scale', '142.89', '-0.6017', '71.45', '26.64', '72.56'),
    ('V4_Complete', '152.25', '-0.6849', '93.50', '22.47', '57.15')
]
for i, row in enumerate(ablation_data, 1):
    cells = table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('图1：N3期EEG信号时域对比（Epoch 852-854）')
run.bold = True

# 添加消融实验图片
img_path = '05_处理结果/消融实验/ablation_visual_comparison.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('从图1可以清晰看到：')
doc.add_paragraph('• 左列：原始信号（蓝）与V4_Complete输出（红）对比，CC≈-0.58')
doc.add_paragraph('• 中列：将去噪信号反转后，与原始信号高度吻合（验证了相位反转假设）')
doc.add_paragraph('• 右列：功率谱密度对比，Delta频段能量确实得到保留')

doc.add_heading('3.2 ASR/RAW/V4三者信号对比实验', level=2)

p = doc.add_paragraph()
p.add_run('为了深入分析深度学习去噪模型与传统ASR算法的差异，本研究基于V4最优去噪模型开展了RAW/ASR/去噪信号的三者对比实验。')

p = doc.add_paragraph()
p.add