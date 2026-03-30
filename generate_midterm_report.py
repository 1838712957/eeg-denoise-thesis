"""
生成毕业设计中期检查报告（Word格式）
包含必要的图表和解释
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent

def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_row(table, cells_text, bold_list=None, bg_color=None):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = str(text)
        paragraph = cell.paragraphs[0]
        if bold_list and i in bold_list:
            for run in paragraph.runs:
                run.font.bold = True
        if bg_color:
            set_cell_shading(cell, bg_color)

def create_report():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ================= 标题 =================
    title = doc.add_heading('', level=0)
    run = title.add_run('毕业设计中期检查报告')
    run.font.size = Pt(22)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ================= 一、课题简介与前期目标回顾 =================
    doc.add_heading('一、课题简介与前期目标回顾', level=1)
    
    doc.add_heading('1.1 研究背景', level=2)
    p = doc.add_paragraph()
    p.add_run('脑电图（EEG）信号去噪在睡眠分期中具有至关重要的作用。然而，传统去噪评估方法存在严重局限性：')
    
    # 添加要点
    bullet_points = [
        '传统指标的欺骗性：仅关注RRMSE、CC等数学指标，无法反映波形的生物学语义',
        '临床特征的丢失：去噪后的信号可能"数学上干净"但"临床上无用"',
        '下游任务性能下降：过度平滑导致关键生理特征（如N3期Delta慢波）被误判为噪声'
    ]
    for point in bullet_points:
        p = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('1.2 原定目标', level=2)
    p = doc.add_paragraph()
    p.add_run('设计并实现一个基于深度学习的睡眠伪差矫正系统，要求：')
    
    goals = [
        '有效去除EEG信号中的噪声干扰',
        '尽可能保护下游临床分期（如N3期Delta慢波）的形态特征',
        '建立多维评估体系，验证去噪效果的临床有效性'
    ]
    for i, goal in enumerate(goals, 1):
        p = doc.add_paragraph(f'{i}. {goal}')
    
    # ================= 二、目前已完成的工作 =================
    doc.add_heading('二、目前已完成的工作', level=1)
    
    doc.add_heading('2.1 数据流水线与基座模型搭建', level=2)
    
    # 数据处理流程图（用文本框表示）
    p = doc.add_paragraph()
    p.add_run('数据处理流程：').bold = True
    doc.add_paragraph('原始EDF文件 → 通道选择(EEG) → 重采样(100Hz) → 分段(30s epoch) → 标签对齐')
    
    p = doc.add_paragraph()
    p.add_run('• 完成了基于Sleep-EDF和DREAMS数据集的数据预处理、重采样与对齐\n'
              '• 实现了自动化的标签解析和epoch分割\n'
              '• 建立了标准化的数据加载接口')
    
    # 裁判模型表格
    doc.add_heading('裁判模型部署', level=3)
    p = doc.add_paragraph()
    p.add_run('成功部署了').bold = False
    run = p.add_run('DeepSleepNet')
    run.bold = True
    p.add_run('作为下游任务的"裁判模型"：')
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ['组件', '说明'], bold_list=[0])
    add_table_row(table, ['输入', '单通道EEG，30s epoch，100Hz采样率'])
    add_table_row(table, ['架构', 'CNN + BiLSTM'])
    add_table_row(table, ['输出', '5类睡眠分期（W, N1, N2, N3, REM）'])
    add_table_row(table, ['验证准确率', '~75%（原始信号）'])
    
    # ================= 2.2 多版本去噪模型 =================
    doc.add_heading('2.2 多版本去噪模型的迭代与消融实验', level=2)
    
    doc.add_heading('模型架构演进', level=3)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ['模型版本', '架构特点', '设计意图'], bold_list=[0, 1, 2])
    add_table_row(table, ['Baseline', '纯1D CNN，无残差、无注意力', '基准对照'])
    add_table_row(table, ['V4_wo_SE', '多尺度残差 + 移除SE注意力', '验证注意力机制的作用'])
    add_table_row(table, ['V4_Single_Scale', '单一尺度(kernel=3) + SE注意力', '验证多尺度的必要性'])
    add_table_row(table, ['V4_Complete', '多尺度残差 + SE注意力', '完整架构'])
    
    # ================= 2.3 核心实验发现 =================
    doc.add_heading('2.3 核心实验发现（高光部分）', level=2)
    
    doc.add_heading('消融实验结果', level=3)
    p = doc.add_paragraph()
    p.add_run('表1：模型变体对比表').bold = True
    
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'
    add_table_row(table, ['模型', 'RRMSE(%)', 'CC', 'Delta保持(%)', '准确率(%)', 'N3召回(%)'], 
                  bold_list=[0], bg_color='D9D9D9')
    add_table_row(table, ['Original (基准)', '-', '-', '-', '11.79', '19.31'])
    add_table_row(table, ['Baseline', '150.30', '-0.6929', '68.52', '28.13', '79.09'])
    add_table_row(table, ['V4_wo_SE', '145.51', '-0.6165', '64.62', '21.81', '62.62'])
    add_table_row(table, ['V4_Single_Scale', '142.89', '-0.6017', '71.45', '26.64', '72.56'])
    add_table_row(table, ['V4_Complete', '152.25', '-0.6849', '93.50', '22.47', '57.15'])
    
    # 添加图片
    doc.add_heading('时域波形可视化分析', level=3)
    p = doc.add_paragraph()
    p.add_run('图1：N3期EEG信号时域对比（Epoch 852-854）').bold = True
    
    img_path = PROJECT_ROOT / "05_处理结果" / "消融实验" / "ablation_visual_comparison.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(6.5))
        p = doc.add_paragraph('左列：原始信号（蓝）与V4_Complete输出（红）对比，CC≈-0.58\n'
                             '中列：将去噪信号反转后，与原始信号高度吻合（验证了相位反转假设）\n'
                             '右列：功率谱密度对比，Delta频段能量确实得到保留')
    else:
        doc.add_paragraph('[图片文件不存在]')
    
    # 发现分析
    doc.add_heading('发现A：频域保真与时域失真的冲突', level=3)
    p = doc.add_paragraph()
    p.add_run('关键发现：').bold = True
    p.add_run('V4_Complete模型保留了高达')
    run = p.add_run('93.5%')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run('的Delta频段能量，但时域相关系数（CC）为')
    run = p.add_run('-0.68（负值！）')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run('这一反常现象揭示了极具学术价值的问题：')
    
    # 文本框
    p = doc.add_paragraph()
    p.add_run('频域能量保持率高 ≠ 时域波形保真\n\n'
              'Delta能量93.5%保持 → 频域"看起来很好"\n'
              'CC = -0.68 → 时域"完全反转"')
    
    doc.add_heading('发现B：信号相位反转问题', level=3)
    p = doc.add_paragraph()
    p.add_run('从时域波形可视化图中可以清晰看到：')
    bullet_points2 = [
        '左列：原始信号（蓝）与V4_Complete输出（红）对比，CC≈-0.58',
        '中列：将去噪信号反转后，与原始信号高度吻合（验证了相位反转假设）',
        '右列：功率谱密度对比，Delta频段能量确实得到保留'
    ]
    for point in bullet_points2:
        p = doc.add_paragraph(point, style='List Bullet')
    
    # 机理解析表格
    p = doc.add_paragraph()
    p.add_run('机理解析：').bold = True
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ['现象', '原因分析'], bold_list=[0], bg_color='D9D9D9')
    add_table_row(table, ['CC为负值', '模型输出信号相位反转（上下颠倒）'])
    add_table_row(table, ['RRMSE > 140%', '尺度坍塌，输出幅度异常放大'])
    add_table_row(table, ['Delta能量保持高', '频域能量对相位不敏感'])
    
    doc.add_heading('发现C：注意力的副作用', level=3)
    p = doc.add_paragraph()
    p.add_run('Baseline模型反而表现更好：').bold = True
    
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    add_table_row(table, ['指标', 'Baseline', 'V4_Complete', '差异'], bold_list=[0], bg_color='D9D9D9')
    add_table_row(table, ['N3召回率', '79.09%', '57.15%', '+21.94pp'])
    add_table_row(table, ['准确率', '28.13%', '22.47%', '+5.66pp'])
    
    p = doc.add_paragraph()
    p.add_run('结论：').bold = True
    p.add_run('复杂的SE注意力机制和多尺度大卷积核，反而加剧了形态学的畸变。')
    
    # ================= 三、存在的主要问题与难点 =================
    doc.add_heading('三、存在的主要问题与难点', level=1)
    
    doc.add_heading('3.1 指标欺骗性问题', level=2)
    p = doc.add_paragraph()
    p.add_run('问题描述：').bold = True
    p.add_run('传统的去噪评价指标（MSE、RRMSE）在面对医疗生理信号时，无法有效反映波形的生物学语义。')
    p = doc.add_paragraph()
    p.add_run('具体表现：').bold = True
    bullet3 = [
        'RRMSE显示去噪效果"良好"，但下游任务准确率下降',
        'CC为负值，说明信号被反转，但频域能量指标无法捕捉这一致命问题'
    ]
    for point in bullet3:
        p = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('3.2 复杂网络的副作用', level=2)
    p = doc.add_paragraph()
    p.add_run('问题描述：').bold = True
    p.add_run('引入SE注意力机制和多尺度大卷积核，反而加剧了形态学的畸变。')
    p = doc.add_paragraph()
    p.add_run('实验证据：').bold = True
    bullet4 = [
        'Baseline（最简单架构）在N3召回率上表现最好',
        'V4_Complete（最复杂架构）在准确率上表现最差'
    ]
    for point in bullet4:
        p = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('3.3 归一化与反归一化缺陷', level=2)
    p = doc.add_paragraph()
    p.add_run('问题描述：').bold = True
    p.add_run('模型训练时的归一化策略可能导致推理时的尺度坍塌和相位反转。')
    p = doc.add_paragraph()
    p.add_run('具体表现：').bold = True
    bullet5 = [
        '输出信号幅度异常放大（RRMSE > 140%）',
        '相位完全反转（CC < 0）'
    ]
    for point in bullet5:
        p = doc.add_paragraph(point, style='List Bullet')
    
    # ================= 四、下一步工作计划 =================
    doc.add_heading('四、下一步工作计划', level=1)
    
    doc.add_heading('4.1 第1-2周：修复与完善算法层面', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    add_table_row(table, ['任务', '具体措施', '预期成果'], bold_list=[0], bg_color='D9D9D9')
    add_table_row(table, ['修复相位反转', '检查反归一化代码，添加相位约束损失', 'CC转正'])
    add_table_row(table, ['修复尺度坍塌', '调整输出层归一化策略', 'RRMSE < 50%'])
    add_table_row(table, ['架构优化', '以Baseline为基础进行改进', '保持简单架构的优势'])
    
    doc.add_heading('4.2 第3-4周：系统设计与实现', level=2)
    p = doc.add_paragraph()
    p.add_run('基于').bold = False
    run = p.add_run('Streamlit')
    run.bold = True
    p.add_run('框架开发轻量级可视化交互网页：')
    bullet6 = [
        'EEG数据上传（支持EDF格式）',
        '去噪波形对比展示（时域+频域）',
        '在线睡眠分期预测',
        'Grad-CAM可解释性分析'
    ]
    for point in bullet6:
        p = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('4.3 第5-6周：论文撰写与图表精修', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ['内容', '要求'], bold_list=[0], bg_color='D9D9D9')
    add_table_row(table, ['PSD频谱图', '高清重绘，标注Delta/Theta/Sigma频段'])
    add_table_row(table, ['Grad-CAM热力图', '展示模型关注区域'])
    add_table_row(table, ['时域对比图', '突出相位反转问题'])
    add_table_row(table, ['消融实验表', '规范学术论文格式'])
    
    # ================= 五、阶段性成果清单 =================
    doc.add_heading('五、阶段性成果清单', level=1)
    
    doc.add_heading('5.1 代码成果', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ['模块', '功能'], bold_list=[0], bg_color='D9D9D9')
    add_table_row(table, ['数据预处理', 'EDF文件读取、重采样、分段'])
    add_table_row(table, ['模型训练', '去噪模型训练脚本'])
    add_table_row(table, ['推理应用', '在线推理引擎'])
    add_table_row(table, ['分期分析', '睡眠分期对比分析'])
    add_table_row(table, ['消融实验', '模型变体对比实验'])
    add_table_row(table, ['可视化', '时域波形对比可视化'])
    
    doc.add_heading('5.2 模型成果', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    add_table_row(table, ['模型', '说明'], bold_list=[0], bg_color='D9D9D9')
    add_table_row(table, ['Baseline.h5', '基础CNN模型'])
    add_table_row(table, ['V4_wo_SE.h5', '移除SE注意力'])
    add_table_row(table, ['V4_Single_Scale.h5', '单一尺度'])
    add_table_row(table, ['V4_Complete.h5', '完整V4模型'])
    add_table_row(table, ['DeepSleepNet裁判模型.h5', '下游分期模型'])
    
    # ================= 六、参考文献 =================
    doc.add_heading('六、参考文献', level=1)
    refs = [
        '[1] Supratak A, Dong H, Wu C, et al. DeepSleepNet: A model for automatic sleep stage scoring based on raw single-channel EEG. 2017.',
        '[2] Mullen T, Kothe C, Chi Y M, et al. Real-time modeling and 3D visualization of source dynamics. 2012.',
        '[3] Delorme A, Makeig S. EEGLAB: an open source toolbox for analysis of single-trial EEG dynamics. 2004.',
        '[4] Hu J, Shen L, Sun G. Squeeze-and-excitation networks. 2018.',
        '[5] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition. 2016.'
    ]
    for ref in refs:
        doc.add_paragraph(ref)
    
    # ================= 七、致谢 =================
    doc.add_heading('七、致谢', level=1)
    p = doc.add_paragraph()
    p.add_run('感谢指导教师的悉心指导，感谢实验室同学的帮助与支持。特别感谢开源社区提供的EEGLAB、MNE-Python等工具包，为本研究的顺利开展提供了重要支持。')
    
    # ================= 页脚 =================
    p = doc.add_paragraph()
    p.add_run('\n\n报告日期：2026年3月26日\n')
    p.add_run('完成进度：约70%')
    
    return doc


if __name__ == "__main__":
    print("正在生成Word格式中期报告...")
    doc = create_report()
    output_path = PROJECT_ROOT / "01_论文文档" / "中期报告_生成版.docx"
    doc.save(str(output_path))
    print(f"报告已生成: {output_path}")
    print("请在Word中打开并检查格式")
