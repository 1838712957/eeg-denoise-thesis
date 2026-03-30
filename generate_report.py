"""
生成中期报告Word文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_data(doc, headers, data, title=None):
    """添加带数据的表格"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'D9E2F3')
    
    # 数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    doc.add_paragraph()
    return table

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_heading('毕业设计中期报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 基本信息
doc.add_heading('一、课题基本信息', level=1)
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ('课题名称', '基于深度学习的脑电信号去噪方法研究'),
    ('学生姓名', '[填写姓名]'),
    ('指导教师', '[填写导师]'),
    ('完成进度', '约70%')
]
for i, (label, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value

doc.add_paragraph()

# 研究背景与意义
doc.add_heading('二、研究背景与意义', level=1)
doc.add_heading('2.1 研究背景', level=2)

p = doc.add_paragraph()
p.add_run('脑电图（EEG）是监测大脑活动的重要工具，广泛应用于睡眠分期、癫痫诊断、脑机接口等领域。然而，原始EEG信号极易受到各种噪声干扰，包括：')

noises = [
    '工频干扰：50/60Hz电源线干扰',
    '肌电伪迹（EMG）：肌肉活动产生的高频噪声',
    '眼电伪迹（EOG）：眼球运动产生的低频干扰',
    '基线漂移：电极接触不良引起的低频波动'
]
for noise in noises:
    doc.add_paragraph(noise, style='List Bullet')

p = doc.add_paragraph()
p.add_run('这些噪声严重影响EEG信号的质量和后续分析的准确性。')

doc.add_heading('2.2 研究意义', level=2)
p = doc.add_paragraph()
p.add_run('开发高效的EEG去噪算法对于提高睡眠分期准确率、辅助临床诊断具有重要意义。传统方法如ASR（Artifact Subspace Reconstruction）存在参数敏感、计算复杂等问题。本研究探索基于深度学习的端到端去噪方法，旨在：')

goals = [
    '提高去噪效率和自动化程度',
    '保留关键生理特征',
    '改善下游任务（如睡眠分期）的性能'
]
for i, goal in enumerate(goals, 1):
    doc.add_paragraph(f'{i}. {goal}')

# 研究内容与方法
doc.add_heading('三、研究内容与方法', level=1)
doc.add_heading('3.1 技术路线', level=2)

p = doc.add_paragraph()
p.add_run('原始EEG数据 → 预处理 → 深度学习去噪模型 → 质量评估 → 下游任务验证')

doc.add_heading('3.2 去噪模型架构', level=2)
p = doc.add_paragraph()
p.add_run('本研究采用基于1D-ResNet的去噪网络，主要特点：')

features = [
    '多尺度残差块：并行使用3×3、5×5、7×7卷积核捕获不同尺度特征',
    'SE注意力机制：自适应调整通道权重，增强特征表达',
    '端到端训练：直接从噪声信号学习到干净信号的映射'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('3.3 评估方法', level=2)
p = doc.add_paragraph()
p.add_run('采用多维度评估体系：')

eval_methods = [
    '信号质量指标：噪声抑制比（NRR）、相关系数（CC）',
    '下游任务验证：使用DeepSleepNet进行睡眠分期，对比去噪前后的分期准确率',
    '频域能量分析：计算各频段能量损失，评估生理特征保留情况'
]
for method in eval_methods:
    doc.add_paragraph(method, style='List Bullet')

# 已完成工作
doc.add_heading('四、已完成工作', level=1)
doc.add_heading('4.1 数据准备', level=2)

headers = ['数据集', '来源', '样本数', '用途']
data = [
    ['Sleep EDF', 'PhysioNet', '7名受试者', '模型训练与测试'],
    ['DREAMS', '比利时大学', '20名受试者', '扩展验证']
]
add_table_with_data(doc, headers, data, '表1 数据集概况')

doc.add_heading('4.2 模型训练', level=2)
p = doc.add_paragraph()
p.add_run('已完成去噪模型v1、v2版本的训练，当前最优模型为V4版本。')

doc.add_heading('4.3 实验结果', level=2)

# 去噪效果对比表
headers = ['受试者', 'NRR_ASR(%)', 'NRR_Ours(%)', '改进(%)']
data = [
    ['SC4001E0', '63.62', '89.43', '+25.81'],
    ['SC4002E0', '53.76', '82.29', '+28.53'],
    ['SC4011E0', '88.34', '95.63', '+7.29'],
    ['SC4012E0', '83.66', '94.81', '+11.15'],
    ['SC4021E0', '85.84', '94.30', '+8.46'],
    ['SC4022E0', '78.79', '96.00', '+17.21'],
    ['SC4031E0', '57.21', '78.71', '+21.50'],
    ['平均', '72.89', '90.17', '+17.28']
]
add_table_with_data(doc, headers, data, '表2 去噪效果对比（NRR指标）')

p = doc.add_paragraph()
run = p.add_run('结论：')
run.bold = True
p.add_run('本方法在噪声抑制比指标上显著优于传统ASR方法，平均提升17.28个百分点。')

# 睡眠分期对比表
headers = ['受试者', 'Raw(%)', 'ASR(%)', 'Ours(%)']
data = [
    ['SC4001', '86.83', '71.92', '70.94'],
    ['SC4002', '89.71', '85.01', '67.73'],
    ['SC4011', '62.56', '50.18', '62.74'],
    ['SC4012', '70.37', '56.04', '65.66'],
    ['SC4022', '76.15', '50.71', '74.52'],
    ['SC4031', '78.30', '71.88', '77.70'],
    ['平均', '77.32', '64.29', '69.72']
]
add_table_with_data(doc, headers, data, '表3 睡眠分期准确率对比')

# 关键发现
doc.add_heading('4.3.3 关键发现：特征过度平滑问题', level=3)

p = doc.add_paragraph()
p.add_run('通过深入分析，发现了一个重要问题：')

headers = ['指标', '原始信号', '去噪后信号', '变化']
data = [
    ['整体分期准确率', '74.30%', '54.01%', '-20.29pp'],
    ['N1期精确率', '35.77%', '12.22%', '-23.56pp'],
    ['N2期精确率', '87.22%', '82.41%', '-4.81pp'],
    ['N3期精确率', '65.55%', '60.87%', '-4.68pp']
]
add_table_with_data(doc, headers, data, '表4 去噪前后分期性能对比')

# 频域能量损失
headers = ['频段', 'N1期能量损失', 'N2期能量损失', 'N3期能量损失']
data = [
    ['Delta (0.5-4Hz)', '-4.38 dB', '-12.52 dB', '-23.21 dB'],
    ['Theta (4-8Hz)', '-21.63 dB', '-22.21 dB', '-22.35 dB'],
    ['Sigma (11-16Hz)', '-60.38 dB', '-52.34 dB', '-66.04 dB']
]
add_table_with_data(doc, headers, data, '表5 频域能量损失分析')

p = doc.add_paragraph()
run = p.add_run('核心发现：')
run.bold = True
p.add_run('深度学习去噪模型虽然有效抑制了噪声，但存在')
run = p.add_run('过度平滑')
run.bold = True
run.font.color.rgb = None
p.add_run('问题，导致关键生理特征（特别是Delta慢波和Sigma纺锤波）被误判为噪声而消除，严重影响睡眠分期准确率。')

# 问题分析与改进方向
doc.add_heading('五、问题分析与改进方向', level=1)
doc.add_heading('5.1 问题诊断', level=2)

problems = [
    '感受野过大：1D-ResNet的大卷积核可能将低频生理信号视为噪声',
    '损失函数单一：仅使用MSE损失，未考虑频域特征保留',
    '训练数据偏差：合成噪声与真实噪声分布存在差异'
]
for i, problem in enumerate(problems, 1):
    doc.add_paragraph(f'{i}. {problem}')

doc.add_heading('5.2 改进方案', level=2)
headers = ['问题', '改进措施', '预期效果']
data = [
    ['感受野过大', '引入多尺度注意力，自适应调整感受野', '保留多尺度生理特征'],
    ['损失函数单一', '添加频域损失、感知损失', '保持频谱结构'],
    ['训练数据偏差', '使用真实噪声数据增强', '提高泛化能力']
]
add_table_with_data(doc, headers, data, '表6 改进方案')

# 后续工作计划
doc.add_heading('六、后续工作计划', level=1)
headers = ['时间', '工作内容', '预期成果']
data = [
    ['第9-10周', '改进损失函数，引入频域约束', '减少能量损失'],
    ['第11-12周', '优化网络结构，添加注意力机制', '提高分期准确率'],
    ['第13-14周', '扩展验证，完善实验对比', '论文实验部分'],
    ['第15-16周', '撰写论文，准备答辩', '完成毕业论文']
]
add_table_with_data(doc, headers, data, '表7 后续工作计划')

# 参考文献
doc.add_heading('七、参考文献', level=1)
refs = [
    '[1] Supratak A, et al. DeepSleepNet: A model for automatic sleep stage scoring based on raw single-channel EEG. IEEE BIBM, 2017.',
    '[2] Mullen T, et al. Real-time modeling and 3D visualization of source dynamics and connectivity using the BCILAB and OpenViBE platforms. Brain-Computer Interface, 2013.',
    '[3] Delorme A, Makeig S. EEGLAB: an open source toolbox for analysis of single-trial EEG dynamics. Journal of Neuroscience Methods, 2004.'
]
for ref in refs:
    doc.add_paragraph(ref)

# 致谢
doc.add_heading('八、致谢', level=1)
p = doc.add_paragraph()
p.add_run('感谢指导教师的悉心指导，感谢实验室同学的帮助与支持。')

# 日期
p = doc.add_paragraph()
p.add_run('\n报告日期：2026年3月25日')

# 保存文档
output_path = '01_论文文档/中期报告.docx'
doc.save(output_path)
print(f'中期报告已生成: {output_path}')